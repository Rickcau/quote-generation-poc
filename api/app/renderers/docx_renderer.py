"""DOCX renderer that builds styled documents directly with python-docx,
matching the HTML preview styling as closely as possible."""

from io import BytesIO
from decimal import Decimal
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_shading(cell, hex_color: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color.lstrip("#")}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)


def _fmt(value) -> str:
    """Format a decimal/float value as currency string."""
    if isinstance(value, Decimal):
        return f"${value:,.2f}"
    return f"${float(value):,.2f}"


def render_docx(html_string: str, quote_data: dict | None = None, style_config: dict | None = None) -> bytes:
    """Render a DOCX document from quote data.

    If quote_data is provided, builds a styled document directly.
    If only html_string is provided, falls back to basic htmldocx conversion.
    """
    if quote_data is None:
        return _fallback_render(html_string)

    config = style_config or {}
    primary_color = config.get("color_scheme", "#1a365d")
    font_family = config.get("font_family", "Calibri")
    primary_rgb = _hex_to_rgb(primary_color)

    doc = Document()

    # -- Configure default style --
    style = doc.styles['Normal']
    style.font.name = font_family
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = primary_rgb
        hs.font.name = font_family

    # -- Set page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    q = quote_data
    customer = q.get('customer', {})

    # -- Header --
    doc.add_heading(q.get('quote_number', ''), level=1)

    p = doc.add_paragraph()
    run = p.add_run('Contoso Environmental Services')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run('\n123 Industrial Parkway, Houston, TX 77001')

    doc.add_paragraph('─' * 60)

    # Prepared for block
    info = doc.add_paragraph()
    info.add_run('Prepared for: ').bold = True
    info.add_run(customer.get('company_name', ''))
    info.add_run('  ')
    info.add_run('Contact: ').bold = True
    contact = customer.get('contact_name', '') or ''
    email = customer.get('contact_email', '') or ''
    info.add_run(f'{contact} ({email})')

    info2 = doc.add_paragraph()
    info2.add_run('Date: ').bold = True
    info2.add_run(str(q.get('quote_date', '')))
    info2.add_run('  ')
    info2.add_run('Valid Until: ').bold = True
    info2.add_run(str(q.get('valid_until', '')))
    info2.add_run('  ')
    info2.add_run('Status: ').bold = True
    info2.add_run(q.get('status', ''))

    # -- Quote Summary --
    doc.add_heading('Quote Summary', level=2)

    line_items = q.get('line_items', [])
    categories = list({item.get('service_category', '') for item in line_items})
    categories.sort()

    p = doc.add_paragraph()
    p.add_run(f'This quote covers environmental services for ')
    run = p.add_run(customer.get('company_name', ''))
    run.bold = True
    p.add_run(f' in the ')
    run = p.add_run(customer.get('industry', ''))
    run.bold = True
    p.add_run(f' sector. The scope includes {len(line_items)} service items across the following categories:')

    for cat in categories:
        doc.add_paragraph(cat, style='List Bullet')

    # -- Services Table --
    doc.add_heading('Services', level=2)

    headers = ['Category', 'Description', 'Qty', 'Unit', 'Unit Price', 'Total']
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        run.font.name = font_family
        _set_cell_shading(cell, primary_color)

    # Data rows
    for idx, item in enumerate(line_items):
        row = table.add_row()
        values = [
            item.get('service_category', ''),
            item.get('service_description', ''),
            str(item.get('quantity', '')),
            item.get('unit', ''),
            _fmt(item.get('unit_price', 0)),
            _fmt(item.get('line_total', 0)),
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            run.font.name = font_family
            # Right-align numeric columns
            if i in (2, 4, 5):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Alternate row shading
            if idx % 2 == 1:
                _set_cell_shading(cell, 'f5f5f5')

    # Set column widths
    widths = [Inches(1.2), Inches(2.4), Inches(0.5), Inches(0.5), Inches(0.9), Inches(0.9)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    # Totals
    doc.add_paragraph()
    tax_rate = q.get('tax_rate', 0)
    if isinstance(tax_rate, Decimal):
        tax_pct = f"{float(tax_rate) * 100:.2f}"
    else:
        tax_pct = f"{float(tax_rate) * 100:.2f}"

    totals = doc.add_paragraph()
    totals.add_run('Subtotal: ').bold = True
    totals.add_run(_fmt(q.get('subtotal', 0)))
    totals.add_run('\n')
    totals.add_run(f'Tax ({tax_pct}%): ').bold = True
    totals.add_run(_fmt(q.get('tax_amount', 0)))
    totals.add_run('\n')
    totals.add_run('Total: ').bold = True
    run = totals.add_run(_fmt(q.get('total', 0)))
    run.bold = True
    run.font.size = Pt(13)

    # -- Regulatory Notes --
    reg_notes = q.get('regulatory_notes')
    if reg_notes:
        doc.add_heading('Regulatory & Compliance Notes', level=2)
        doc.add_paragraph(reg_notes)

    # -- Terms --
    doc.add_heading('Terms & Conditions', level=2)
    valid_until = str(q.get('valid_until', ''))
    terms = [
        f'This quote is valid until {valid_until}.',
        'Payment is due within 30 days of service completion.',
        'All services are performed in accordance with applicable federal, state, and local regulations.',
        'Pricing is based on estimated quantities; actual charges may vary based on field conditions.',
        'Cancellation within 48 hours of scheduled service may incur a mobilization fee.',
    ]
    for i, term in enumerate(terms, 1):
        doc.add_paragraph(f'{i}. {term}')

    # -- Signature Block --
    doc.add_heading('Authorization', level=2)
    doc.add_paragraph(
        'By signing below, the client authorizes Contoso Environmental Services '
        'to proceed with the services described in this quote.'
    )
    doc.add_paragraph()

    sig_table = doc.add_table(rows=5, cols=2)
    sig_table.style = 'Table Grid'
    sig_labels = ['Client Signature:', 'Printed Name:', 'Date:', 'Contoso Representative:', 'Date:']
    for i, label in enumerate(sig_labels):
        cell = sig_table.rows[i].cells[0]
        cell.text = ''
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        sig_table.rows[i].cells[1].text = ''

    # -- Save --
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _fallback_render(html_string: str) -> bytes:
    """Fallback: basic htmldocx conversion when no quote_data provided."""
    import re
    from htmldocx import HtmlToDocx

    html_string = re.sub(r'<li>\s*<p>(.*?)</p>\s*</li>', r'<li>\1</li>', html_string, flags=re.DOTALL)
    doc = Document()
    parser = HtmlToDocx()
    parser.add_html_to_document(html_string, doc)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
